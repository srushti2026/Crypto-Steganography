#!/usr/bin/env python3
"""
Convert SETUP.md to Word document with proper formatting
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    # This is a simplified hyperlink - Word documents need proper hyperlink formatting
    run = paragraph.add_run(text)
    run.font.color.rgb = None  # Use default hyperlink color
    return run

def convert_markdown_to_word(md_file_path, output_path):
    """Convert markdown file to Word document"""
    
    # Read the markdown content
    with open(md_file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Create a new Word document
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Split content into lines
    lines = content.split('\n')
    
    # Process each line
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # Empty line - add paragraph break
            doc.add_paragraph()
            i += 1
            continue
        
        # Title (# heading)
        if line.startswith('# '):
            title = line[2:].strip()
            # Remove emoji and clean title
            title = re.sub(r'[🛡️🌟✨🎯🕵️🔧🏗📋🚀🌐⚙️📚🔧🔒📞📄]', '', title).strip()
            
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Main headings (## heading)
        elif line.startswith('## '):
            title = line[3:].strip()
            # Remove emoji and clean title
            title = re.sub(r'[🛡️🌟✨🎯🕵️🔧🏗📋🚀🌐⚙️📚🔧🔒📞📄]', '', title).strip()
            doc.add_heading(title, level=2)
            
        # Sub headings (### heading)
        elif line.startswith('### '):
            title = line[4:].strip()
            # Remove emoji and clean title  
            title = re.sub(r'[🛡️🌟✨🎯🕵️🔧🏗📋🚀🌐⚙️📚🔧🔒📞📄]', '', title).strip()
            doc.add_heading(title, level=3)
            
        # Code blocks
        elif line.startswith('```'):
            # Start of code block
            i += 1
            code_content = []
            
            # Collect code block content
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_content.append(lines[i])
                i += 1
            
            # Add code block as formatted text
            if code_content:
                code_para = doc.add_paragraph()
                code_run = code_para.add_run('\n'.join(code_content))
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
                
                # Set paragraph formatting for code
                code_para.paragraph_format.left_indent = Inches(0.5)
                code_para.paragraph_format.space_before = Pt(6)
                code_para.paragraph_format.space_after = Pt(6)
        
        # Tables (simplified - detect table headers)
        elif '|' in line and line.count('|') >= 2:
            # Table detected
            table_lines = [line]
            i += 1
            
            # Collect table content
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            
            # Process table
            if len(table_lines) >= 2:
                # Parse table structure
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                
                # Skip separator line
                if len(table_lines) > 2:
                    rows = []
                    for table_line in table_lines[2:]:
                        if table_line.strip():
                            row = [cell.strip() for cell in table_line.split('|')[1:-1]]
                            if len(row) == len(headers):
                                rows.append(row)
                    
                    # Create Word table
                    if rows:
                        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                        table.style = 'Light Grid Accent 1'
                        
                        # Add headers
                        header_row = table.rows[0]
                        for j, header in enumerate(headers):
                            header_row.cells[j].text = header
                            
                        # Add data rows
                        for row_idx, row in enumerate(rows):
                            table_row = table.rows[row_idx + 1]
                            for col_idx, cell_data in enumerate(row):
                                table_row.cells[col_idx].text = cell_data
            
            i -= 1  # Adjust counter
            
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            # Remove emoji from bullet points
            bullet_text = re.sub(r'[🖼️🔐🕵️📦🎯🌐🔒📊]', '', bullet_text).strip()
            doc.add_paragraph(bullet_text, style='List Bullet')
            
        # Numbered lists
        elif re.match(r'^\d+\.', line):
            list_text = re.sub(r'^\d+\.\s*', '', line).strip()
            # Remove emoji from numbered lists
            list_text = re.sub(r'[🛡️🌟✨🎯🕵️🔧🏗📋🚀🌐⚙️📚🔧🔒📞📄]', '', list_text).strip()
            doc.add_paragraph(list_text, style='List Number')
            
        # Regular paragraphs
        else:
            # Clean up emoji and markdown formatting
            clean_line = line
            
            # Remove markdown bold/italic
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_line)
            clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
            
            # Remove inline code backticks
            clean_line = re.sub(r'`([^`]+)`', r'\1', clean_line)
            
            # Remove emoji
            clean_line = re.sub(r'[🛡️🌟✨🎯🕵️🔧🏗📋🚀🌐⚙️📚🔧🔒📞📄💾🎯🔍💻🔑📁📤📥🔧💾]', '', clean_line)
            
            # Handle links - simplified
            clean_line = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean_line)
            
            if clean_line.strip():
                para = doc.add_paragraph(clean_line.strip())
        
        i += 1
    
    # Save the document
    doc.save(output_path)
    print(f"✅ Successfully converted to Word document: {output_path}")

if __name__ == "__main__":
    # Define paths
    md_file = Path("SETUP.md")
    word_file = Path("VeilForge_Setup_Guide.docx")
    
    # Check if markdown file exists
    if not md_file.exists():
        print(f"❌ Error: {md_file} not found!")
        exit(1)
    
    # Convert to Word
    try:
        convert_markdown_to_word(md_file, word_file)
        print(f"📄 Word document created successfully!")
        print(f"📍 Location: {word_file.absolute()}")
    except Exception as e:
        print(f"❌ Error converting to Word: {e}")
        exit(1)